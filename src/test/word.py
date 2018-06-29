# -*- coding: utf-8 -*-
from docx import Document
import os
import fnmatch
import time
# document = Document()
# nm="test.docx"
# paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
# document.save(nm)

document0 = Document()


path='E:\\b_project\\py\\src\\test\\'
path='C:\\Users\\Administrator\\Desktop\\cs111\\'

# try:
#     for path, dirs, files in os.walk(path):
#         for filename in files:
#             if not fnmatch.fnmatch(filename, '*.docx'): continue
#             document1 = Document(path+filename)
#             for r in document1.paragraphs:
#                 document0.add_paragraph(r.text)
# finally:
#    document0.save("out%s.docx"%int(time.time()))

try:
    for path, dirs, files in os.walk(path):
        for filename in files:
            if not fnmatch.fnmatch(filename, '*.docx'): continue
            document1 = Document(path+filename)
            document1.save("out%s.docx" % int(time.time()))
finally:
   pass
